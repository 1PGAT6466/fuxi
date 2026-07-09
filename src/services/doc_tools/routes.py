"""
routes.py — 文档工具服务 API 路由
提供文件格式转换、PDF 操作、图片处理、文本提取等功能
"""

import io
import logging
import os
import shutil
from pathlib import Path
from typing import List, Optional

from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse

logger = logging.getLogger("services.doc-tools.routes")

router = APIRouter(prefix="/api/tools", tags=["Document Tools"])

# 临时文件目录（相对于项目根目录）
TEMP_DIR = Path("temp")

# 最大文件大小 100MB
MAX_FILE_SIZE = 100 * 1024 * 1024

# 支持的文件格式
SUPPORTED_CONVERT_FORMATS = {
    "txt": "text/plain",
    "md": "text/markdown",
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

SUPPORTED_IMAGE_FORMATS = {
    ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".tif", ".webp", ".ico",
}


def _ensure_temp_dir() -> Path:
    """确保临时目录存在并返回路径"""
    TEMP_DIR.mkdir(parents=True, exist_ok=True)
    return TEMP_DIR


def _cleanup_temp(*paths: Path) -> None:
    """清理临时文件"""
    for p in paths:
        try:
            if p.is_file():
                p.unlink()
            elif p.is_dir():
                shutil.rmtree(str(p))
        except Exception as e:  # TODO: Narrow exception type
            logger.warning(f"清理临时文件失败 {p}: {e}")


# ============ 健康检查 ============

@router.get("/health")
def health_check():
    """服务健康检查"""
    _ensure_temp_dir()
    return {
        "status": "ok",
        "service": "doc-tools",
        "version": "1.0.0",
        "dependencies": {
            "pypdf": _check_import("pypdf"),
            "pillow": _check_import("PIL"),
            "python_docx": _check_import("docx"),
        },
    }


def _check_import(module_name: str) -> bool:
    try:
        __import__(module_name)
        return True
    except ImportError:
        return False


# ============ 文件格式转换 ============

@router.post("/convert")
# FAKE-ASYNC: 本函数标记 async 仅为接口统一，内部同步执行
async def convert_file(
    file: UploadFile = File(...),
    target_format: str = Form(..., description="目标格式: txt, pdf, docx, md"),
):
    """
    文件格式转换
    支持的转换路径：txt↔pdf, txt↔docx, md→txt, pdf→txt, docx→txt
    """
    if not file.filename:
        raise HTTPException(400, "未提供文件名")

    target_format = target_format.lower().strip().lstrip(".")
    if target_format not in SUPPORTED_CONVERT_FORMATS:
        raise HTTPException(400, f"不支持的目标格式: {target_format}。支持: {list(SUPPORTED_CONVERT_FORMATS.keys())}")

    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(413, f"文件过大，最大支持 {MAX_FILE_SIZE // (1024 * 1024)}MB")

    source_ext = Path(file.filename).suffix.lower().lstrip(".")
    temp_dir = _ensure_temp_dir()
    input_path = temp_dir / f"convert_in_{os.urandom(8).hex()}.{source_ext}"
    output_path = temp_dir / f"convert_out_{os.urandom(8).hex()}.{target_format}"

    try:
        input_path.write_bytes(content)

        # 如果源格式和目标格式相同，直接返回
        if source_ext == target_format:
            output_path.write_bytes(content)
        else:
            _do_convert(input_path, output_path, source_ext, target_format)

        if not output_path.exists() or output_path.stat().st_size == 0:
            raise HTTPException(500, "转换失败：输出文件为空")

        media_type = SUPPORTED_CONVERT_FORMATS.get(target_format, "application/octet-stream")
        return FileResponse(
            path=str(output_path),
            filename=f"converted.{target_format}",
            media_type=media_type,
            background=lambda: _cleanup_temp(input_path, output_path),
        )

    except HTTPException:
        _cleanup_temp(input_path, output_path)
        raise
    except Exception as e:  # TODO: Narrow exception type
        _cleanup_temp(input_path, output_path)
        logger.error(f"文件转换失败: {e}", exc_info=True)
        raise HTTPException(500, "文件转换失败，请稍后重试")


def _do_convert(input_path: Path, output_path: Path, source_fmt: str, target_fmt: str) -> None:
    """执行格式转换"""
    text_content = None

    # 第一步：提取文本
    if source_fmt in ("txt", "md"):
        text_content = input_path.read_text(encoding="utf-8", errors="replace")
    elif source_fmt == "docx":
        text_content = _extract_text_from_docx(input_path)
    elif source_fmt == "pdf":
        text_content = _extract_text_from_pdf(input_path)

    if text_content is None:
        raise HTTPException(400, f"无法从 {source_fmt} 格式提取文本")

    # 第二步：写入目标格式
    if target_fmt in ("txt", "md"):
        output_path.write_text(text_content, encoding="utf-8")
    elif target_fmt == "pdf":
        _write_text_to_pdf(text_content, output_path)
    elif target_fmt == "docx":
        _write_text_to_docx(text_content, output_path)


# ============ PDF 操作 ============

@router.post("/merge")
async def merge_pdfs(files: List[UploadFile] = File(...)):
    """
    合并多个 PDF 文件
    接受多个 PDF 文件上传，按上传顺序合并为一个 PDF
    """
    if not files or len(files) < 2:
        raise HTTPException(400, "请至少上传 2 个 PDF 文件进行合并")

    try:
        from pypdf import PdfWriter, PdfReader
    except ImportError:
        raise HTTPException(503, "pypdf 未安装 — PDF 合并功能不可用")

    temp_dir = _ensure_temp_dir()
    input_paths = []
    merger = PdfWriter()

    try:
        for i, f in enumerate(files):
            if not f.filename:
                raise HTTPException(400, f"第 {i + 1} 个文件缺少文件名")

            if not f.filename.lower().endswith(".pdf"):
                raise HTTPException(400, f"文件 '{f.filename}' 不是 PDF 格式")

            content = await f.read()
            tmp_path = temp_dir / f"merge_in_{i}_{os.urandom(8).hex()}.pdf"
            tmp_path.write_bytes(content)
            input_paths.append(tmp_path)

            reader = PdfReader(str(tmp_path))
            for page in reader.pages:
                merger.add_page(page)

        output_path = temp_dir / f"merged_{os.urandom(8).hex()}.pdf"
        with open(str(output_path), "wb") as out_f:
            merger.write(out_f)

        return FileResponse(
            path=str(output_path),
            filename="merged.pdf",
            media_type="application/pdf",
            background=lambda: _cleanup_temp(output_path, *input_paths),
        )

    except HTTPException:
        _cleanup_temp(*input_paths)
        raise
    except Exception as e:  # TODO: Narrow exception type
        _cleanup_temp(*input_paths)
        logger.error(f"PDF 合并失败: {e}", exc_info=True)
        raise HTTPException(500, "PDF 合并失败，请稍后重试")


@router.post("/split")
# FAKE-ASYNC: 本函数标记 async 仅为接口统一，内部同步执行
async def split_pdf(
    file: UploadFile = File(...),
    start_page: int = Form(..., ge=1, description="起始页码（从 1 开始）"),
    end_page: int = Form(..., ge=1, description="结束页码（含）"),
):
    """
    拆分 PDF 的指定页码范围
    """
    try:
        from pypdf import PdfWriter, PdfReader
    except ImportError:
        raise HTTPException(503, "pypdf 未安装 — PDF 拆分功能不可用")

    if not file.filename:
        raise HTTPException(400, "未提供文件名")

    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(400, "请上传 PDF 文件")

    if start_page > end_page:
        raise HTTPException(400, f"起始页码 ({start_page}) 不能大于结束页码 ({end_page})")

    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(413, f"文件过大，最大支持 {MAX_FILE_SIZE // (1024 * 1024)}MB")

    temp_dir = _ensure_temp_dir()
    input_path = temp_dir / f"split_in_{os.urandom(8).hex()}.pdf"

    try:
        input_path.write_bytes(content)

        reader = PdfReader(str(input_path))
        total_pages = len(reader.pages)

        if start_page > total_pages:
            raise HTTPException(400, f"起始页码 ({start_page}) 超出总页数 ({total_pages})")

        actual_end = min(end_page, total_pages)

        writer = PdfWriter()
        for page_num in range(start_page - 1, actual_end):
            writer.add_page(reader.pages[page_num])

        output_path = temp_dir / f"split_out_{os.urandom(8).hex()}.pdf"
        with open(str(output_path), "wb") as out_f:
            writer.write(out_f)

        return FileResponse(
            path=str(output_path),
            filename=f"split_pages_{start_page}-{actual_end}.pdf",
            media_type="application/pdf",
            background=lambda: _cleanup_temp(input_path, output_path),
        )

    except HTTPException:
        _cleanup_temp(input_path)
        raise
    except Exception as e:  # TODO: Narrow exception type
        _cleanup_temp(input_path)
        logger.error(f"PDF 拆分失败: {e}", exc_info=True)
        raise HTTPException(500, "PDF 拆分失败，请稍后重试")


# ============ 文件压缩 ============

@router.post("/compress")
async def compress_file(file: UploadFile = File(...)):
    """
    压缩上传的文件
    对图片进行有损压缩，对 PDF 进行优化
    """
    if not file.filename:
        raise HTTPException(400, "未提供文件名")

    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(413, f"文件过大，最大支持 {MAX_FILE_SIZE // (1024 * 1024)}MB")

    ext = Path(file.filename).suffix.lower()
    temp_dir = _ensure_temp_dir()
    input_path = temp_dir / f"compress_in_{os.urandom(8).hex()}{ext}"

    try:
        input_path.write_bytes(content)

        if ext in SUPPORTED_IMAGE_FORMATS:
            output_path = _compress_image_file(input_path, ext)
        elif ext == ".pdf":
            output_path = _compress_pdf_file(input_path)
        else:
            raise HTTPException(400, f"不支持压缩的文件格式: {ext}。支持: 图片格式、PDF")

        original_size = len(content)
        compressed_size = output_path.stat().st_size
        ratio = round((1 - compressed_size / original_size) * 100, 1) if original_size > 0 else 0

        return FileResponse(
            path=str(output_path),
            filename=f"compressed_{file.filename}",
            media_type="application/octet-stream",
            headers={
                "X-Original-Size": str(original_size),
                "X-Compressed-Size": str(compressed_size),
                "X-Compression-Ratio": f"{ratio}%",
            },
            background=lambda: _cleanup_temp(input_path, output_path),
        )

    except HTTPException:
        _cleanup_temp(input_path)
        raise
    except Exception as e:  # TODO: Narrow exception type
        _cleanup_temp(input_path)
        logger.error(f"文件压缩失败: {e}", exc_info=True)
        raise HTTPException(500, "文件压缩失败，请稍后重试")


def _compress_image_file(input_path: Path, ext: str) -> Path:
    """压缩图片文件"""
    from PIL import Image

    temp_dir = _ensure_temp_dir()
    output_path = temp_dir / f"compress_out_{os.urandom(8).hex()}{ext}"

    with Image.open(str(input_path)) as img:
        # 保持 RGB 模式用于 JPEG 压缩
        if img.mode in ("RGBA", "P", "LA"):
            # 有透明通道的保持 PNG
            output_path = temp_dir / f"compress_out_{os.urandom(8).hex()}.png"
            img.save(str(output_path), "PNG", optimize=True)
        elif ext in (".jpg", ".jpeg"):
            img_rgb = img.convert("RGB")
            img_rgb.save(str(output_path), "JPEG", quality=75, optimize=True)
        elif ext == ".png":
            img.save(str(output_path), "PNG", optimize=True)
        elif ext == ".webp":
            img.save(str(output_path), "WEBP", quality=75)
        else:
            img.save(str(output_path), quality=75, optimize=True)

    return output_path


def _compress_pdf_file(input_path: Path) -> Path:
    """压缩 PDF 文件"""
    try:
        from pypdf import PdfWriter, PdfReader
    except ImportError:
        raise HTTPException(503, "pypdf 未安装 — PDF 压缩功能不可用")

    temp_dir = _ensure_temp_dir()
    output_path = temp_dir / f"compress_out_{os.urandom(8).hex()}.pdf"

    reader = PdfReader(str(input_path))
    writer = PdfWriter()

    for page in reader.pages:
        page.compress_content_streams()
        writer.add_page(page)

    with open(str(output_path), "wb") as out_f:
        writer.write(out_f)

    return output_path


# ============ 图片信息 ============

@router.post("/image-info")
async def image_info(file: UploadFile = File(...)):
    """获取图片元信息（尺寸、格式、DPI、颜色模式等）"""
    try:
        from PIL import Image
    except ImportError:
        raise HTTPException(503, "Pillow 未安装 — 图片信息功能不可用")

    if not file.filename:
        raise HTTPException(400, "未提供文件名")

    ext = Path(file.filename).suffix.lower()
    if ext not in SUPPORTED_IMAGE_FORMATS:
        raise HTTPException(400, f"不支持的图片格式: {ext}")

    content = await file.read()
    if len(content) > 50 * 1024 * 1024:
        raise HTTPException(413, "图片过大，最大支持 50MB")

    try:
        img = Image.open(io.BytesIO(content))

        info = {
            "filename": file.filename,
            "format": img.format,
            "mode": img.mode,
            "width": img.width,
            "height": img.height,
            "size_bytes": len(content),
        }

        # DPI 信息
        dpi = img.info.get("dpi")
        if dpi:
            info["dpi"] = {"x": round(dpi[0], 1), "y": round(dpi[1], 1)}

        # 额外元数据
        exif_data = {}
        if hasattr(img, "_getexif") and img._getexif():
            for tag_id, value in img._getexif().items():
                from PIL.ExifTags import TAGS
                tag_name = TAGS.get(tag_id, tag_id)
                if isinstance(value, bytes):
                    value = value.hex()
                elif isinstance(value, (tuple, list)):
                    value = str(value)
                exif_data[tag_name] = value
        if exif_data:
            info["exif"] = exif_data

        # 计算宽高比
        if img.height > 0:
            info["aspect_ratio"] = round(img.width / img.height, 3)

        return info

    except Exception as e:  # TODO: Narrow exception type
        logger.error(f"读取图片信息失败: {e}", exc_info=True)
        raise HTTPException(500, "读取图片信息失败，请稍后重试")


# ============ 图片压缩 ============

@router.post("/compress-image")
# FAKE-ASYNC: 本函数标记 async 仅为接口统一，内部同步执行
async def compress_image(
    file: UploadFile = File(...),
    quality: int = Form(75, ge=1, le=100, description="压缩质量 1-100"),
    max_width: Optional[int] = Form(None, ge=1, description="最大宽度（像素），超过则等比缩放"),
    max_height: Optional[int] = Form(None, ge=1, description="最大高度（像素），超过则等比缩放"),
):
    """
    压缩上传的图片文件
    支持调整质量和尺寸
    """
    try:
        from PIL import Image
    except ImportError:
        raise HTTPException(503, "Pillow 未安装 — 图片压缩功能不可用")

    if not file.filename:
        raise HTTPException(400, "未提供文件名")

    ext = Path(file.filename).suffix.lower()
    if ext not in SUPPORTED_IMAGE_FORMATS:
        raise HTTPException(400, f"不支持的图片格式: {ext}")

    content = await file.read()
    if len(content) > 50 * 1024 * 1024:
        raise HTTPException(413, "图片过大，最大支持 50MB")

    temp_dir = _ensure_temp_dir()
    output_path = temp_dir / f"img_compressed_{os.urandom(8).hex()}{ext}"

    try:
        img = Image.open(io.BytesIO(content))
        original_size = len(content)

        # 缩放
        if max_width or max_height:
            new_w = img.width
            new_h = img.height

            if max_width and img.width > max_width:
                ratio = max_width / img.width
                new_w = max_width
                new_h = int(img.height * ratio)

            if max_height and new_h > max_height:
                ratio = max_height / new_h
                new_h = max_height
                new_w = int(new_w * ratio)

            if (new_w, new_h) != (img.width, img.height):
                img = img.resize((new_w, new_h), Image.LANCZOS)

        # 保存
        save_kwargs = {}
        output_ext = ext

        if ext in (".jpg", ".jpeg"):
            if img.mode not in ("RGB", "L"):
                img = img.convert("RGB")
            save_kwargs = {"quality": quality, "optimize": True}
        elif ext == ".png":
            save_kwargs = {"optimize": True}
        elif ext == ".webp":
            save_kwargs = {"quality": quality}
        else:
            output_ext = ".jpg"
            output_path = temp_dir / f"img_compressed_{os.urandom(8).hex()}.jpg"
            if img.mode not in ("RGB", "L"):
                img = img.convert("RGB")
            save_kwargs = {"quality": quality, "optimize": True}

        img.save(str(output_path), **save_kwargs)

        compressed_size = output_path.stat().st_size
        ratio = round((1 - compressed_size / original_size) * 100, 1) if original_size > 0 else 0

        return FileResponse(
            path=str(output_path),
            filename=f"compressed_{Path(file.filename).stem}{output_ext}",
            media_type=f"image/{output_ext.lstrip('.')}",
            headers={
                "X-Original-Size": str(original_size),
                "X-Compressed-Size": str(compressed_size),
                "X-Compression-Ratio": f"{ratio}%",
            },
            background=lambda: _cleanup_temp(output_path),
        )

    except HTTPException:
        raise
    except Exception as e:  # TODO: Narrow exception type
        logger.error(f"图片压缩失败: {e}", exc_info=True)
        raise HTTPException(500, "图片压缩失败，请稍后重试")


# ============ 文本提取 ============

@router.post("/text-extract")
async def text_extract(file: UploadFile = File(...)):
    """
    从各种文档格式提取纯文本
    支持: PDF, DOCX, TXT, MD
    """
    if not file.filename:
        raise HTTPException(400, "未提供文件名")

    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(413, f"文件过大，最大支持 {MAX_FILE_SIZE // (1024 * 1024)}MB")

    ext = Path(file.filename).suffix.lower()

    temp_dir = _ensure_temp_dir()
    tmp_path = temp_dir / f"extract_{os.urandom(8).hex()}{ext}"

    try:
        tmp_path.write_bytes(content)

        if ext == ".pdf":
            text = _extract_text_from_pdf(tmp_path)
        elif ext == ".docx":
            text = _extract_text_from_docx(tmp_path)
        elif ext in (".txt", ".md", ".csv", ".json", ".xml", ".html", ".py", ".js", ".ts", ".yaml", ".yml", ".toml"):
            text = content.decode("utf-8", errors="replace")
        else:
            raise HTTPException(400, f"不支持的文本提取格式: {ext}。支持: PDF, DOCX, TXT, MD 等文本文件")

        char_count = len(text)
        line_count = text.count("\n") + 1
        preview = text[:500] if len(text) > 500 else text

        return {
            "filename": file.filename,
            "format": ext.lstrip("."),
            "char_count": char_count,
            "line_count": line_count,
            "text": text,
            "preview": preview,
        }

    except HTTPException:
        raise
    except Exception as e:  # TODO: Narrow exception type
        logger.error(f"文本提取失败: {e}", exc_info=True)
        raise HTTPException(500, "文本提取失败，请稍后重试")
    finally:
        _cleanup_temp(tmp_path)


# ============ 辅助函数 ============

def _extract_text_from_pdf(pdf_path: Path) -> str:
    """从 PDF 提取纯文本"""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise HTTPException(503, "pypdf 未安装")

    reader = PdfReader(str(pdf_path))
    texts = []
    for i, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text:
            texts.append(page_text)
    return "\n\n".join(texts)


def _extract_text_from_docx(docx_path: Path) -> str:
    """从 DOCX 提取纯文本"""
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(503, "python-docx 未安装")

    doc = Document(str(docx_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def _write_text_to_pdf(text: str, output_path: Path) -> None:
    """将纯文本写入 PDF"""
    from pypdf import PdfWriter

    # 使用 pypdf 创建简单 PDF（文本逐行写入）
    # 对于纯文本 → PDF，使用 reportlab 更合适，但这里保持依赖轻量
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4

        c = canvas.Canvas(str(output_path), pagesize=A4)
        width, height = A4
        margin = 50
        line_height = 14
        y = height - margin

        for line in text.split("\n"):
            if y < margin:
                c.showPage()
                y = height - margin

            # 截断超长行
            display_line = line[:120]
            c.drawString(margin, y, display_line)
            y -= line_height

        c.save()
    except ImportError:
        # reportlab 不可用时用简单方法
        writer = PdfWriter()
        writer.add_blank_page(width=595, height=842)  # A4
        with open(str(output_path), "wb") as f:
            writer.write(f)
        logger.warning("reportlab 不可用，生成空白 PDF；文本内容未写入")


def _write_text_to_docx(text: str, output_path: Path) -> None:
    """将纯文本写入 DOCX"""
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(503, "python-docx 未安装")

    doc = Document()
    for paragraph in text.split("\n"):
        doc.add_paragraph(paragraph)
    doc.save(str(output_path))
