"""
file_connector.py — 文件系统连接器
=================================
支持从本地文件系统批量读取文件内容。
支持格式：.txt, .md, .json, .csv, .yaml/.yml, .html, .pdf, .docx, .xlsx
"""
import csv
import json
import logging
import os
from pathlib import Path
from typing import Any, Dict, List, Optional

from .base import (
    DataSource,
    ConnectorConfig,
    UnifiedDocument,
    SourceType,
    ConnectorStatus,
    ConnectionError,
    FetchError,
    TransformError,
)

logger = logging.getLogger(__name__)


class FileConnector(DataSource):
    """
    FileConnector — 文件系统知识接入连接器

    批量读取指定目录/文件，抽取文本内容转换为统一文档格式。

    支持的文件格式:
    - 纯文本: .txt, .md, .rst, .log
    - 结构化: .json, .csv, .yaml, .yml
    - 标记: .html, .htm, .xml
    - 二进制文档: .pdf, .docx, .xlsx (需要可选依赖)

    配置示例::

        config = ConnectorConfig(
            name="技术文档目录",
            source_type=SourceType.FILE,
            extra={
                "root_path": "/data/documents",
                "patterns": ["*.md", "*.txt"],
                "recursive": True,
                "max_file_size_mb": 50,
            }
        )
        connector = FileConnector(config)
        docs = await connector.ingest()
    """

    # 支持的文件格式及其处理方法
    _TEXT_EXTENSIONS = {".txt", ".md", ".rst", ".log", ".py", ".js", ".ts",
                        ".java", ".go", ".rs", ".cpp", ".c", ".h", ".css",
                        ".html", ".htm", ".xml", ".yaml", ".yml", ".toml",
                        ".ini", ".cfg", ".conf", ".sh", ".bat", ".sql"}
    _STRUCTURED_EXTENSIONS = {".json", ".csv"}
    _BINARY_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".pptx", ".doc"}

    def __init__(self, config: ConnectorConfig):
        super().__init__(config)
        self._root_path: Optional[Path] = None
        self._files_found: int = 0
        self._files_processed: int = 0
        self._files_failed: int = 0

    async def connect(self) -> bool:
        """
        验证文件系统路径可达。

        Returns:
            bool: 路径有效返回 True

        Raises:
            ConnectionError: 路径不存在或无权限时抛出
        """
        self._set_status(ConnectorStatus.CONNECTING)

        try:
            self._validate_config(["root_path"])
            root = self.config.extra["root_path"]
            self._root_path = Path(root).resolve()

            if not self._root_path.exists():
                raise ConnectionError(
                    self.name,
                    f"路径不存在: {self._root_path}"
                )

            if self._root_path.is_file():
                logger.info(
                    "[%s] 文件路径就绪: %s",
                    self.name, self._root_path.name
                )
            else:
                logger.info(
                    "[%s] 目录路径就绪: %s",
                    self.name, str(self._root_path)
                )

            self._set_status(ConnectorStatus.CONNECTED)
            return True

        except ConnectionError:
            self._set_error(str(ConnectionError))
            raise
        except Exception as e:
            self._set_error(str(e))
            raise ConnectionError(self.name, str(e))

    async def fetch(self, **kwargs) -> List[Dict[str, Any]]:
        """
        扫描文件系统并收集文件路径。

        Args:
            **kwargs: 可选参数
                - root_path: 覆盖配置中的路径
                - patterns: 文件匹配模式列表
                - recursive: 是否递归搜索子目录
                - max_files: 最大文件数

        Returns:
            List[Dict[str, Any]]: 文件信息列表
                [{"path": str, "name": str, "size": int, "ext": str}, ...]

        Raises:
            FetchError: 扫描失败时抛出
        """
        root_path = Path(kwargs.get(
            "root_path",
            self.config.extra["root_path"]
        )).resolve()
        patterns = kwargs.get(
            "patterns",
            self.config.extra.get("patterns", ["*"])
        )
        recursive = kwargs.get(
            "recursive",
            self.config.extra.get("recursive", True)
        )
        max_files = kwargs.get(
            "max_files",
            self.config.extra.get("max_files", 10000)
        )

        try:
            files = []
            max_size = self.config.extra.get(
                "max_file_size_mb", 100
            ) * 1024 * 1024

            if root_path.is_file():
                # 单文件模式
                stat = root_path.stat()
                files.append({
                    "path": str(root_path),
                    "name": root_path.name,
                    "size": stat.st_size,
                    "ext": root_path.suffix.lower(),
                })
            else:
                # 目录模式
                for pattern in patterns:
                    if recursive:
                        matches = root_path.rglob(pattern)
                    else:
                        matches = root_path.glob(pattern)

                    for file_path in matches:
                        if not file_path.is_file():
                            continue
                        if file_path.stat().st_size > max_size:
                            logger.warning(
                                "[%s] 跳过超大文件: %s (%.1f MB)",
                                self.name, file_path.name,
                                file_path.stat().st_size / 1024 / 1024
                            )
                            continue
                        if len(files) >= max_files:
                            break

                        files.append({
                            "path": str(file_path),
                            "name": file_path.name,
                            "size": file_path.stat().st_size,
                            "ext": file_path.suffix.lower(),
                            "mtime": file_path.stat().st_mtime,
                        })

                    if len(files) >= max_files:
                        break

            self._files_found = len(files)
            logger.info(
                "[%s] 扫描完成: 找到 %d 个文件",
                self.name, self._files_found
            )
            return files

        except Exception as e:
            raise FetchError(self.name, str(e))

    async def transform(self, raw_data: List[Dict[str, Any]]) -> List[UnifiedDocument]:
        """
        读取文件内容并转换为统一文档格式。

        根据文件扩展名选择不同的读取策略：
        - 文本文件：直接读取字符内容
        - JSON：解析为结构化内容
        - CSV：逐行读取
        - PDF/DOCX/XLSX：尝试使用第三方库提取文本

        Args:
            raw_data: fetch() 返回的文件信息列表

        Returns:
            List[UnifiedDocument]: 统一文档列表

        Raises:
            TransformError: 读取失败时抛出
        """
        if not raw_data:
            return []

        documents = []
        self._files_processed = 0
        self._files_failed = 0

        for file_info in raw_data:
            try:
                doc = await self._process_file(file_info)
                if doc:
                    documents.append(doc)
                    self._files_processed += 1
                else:
                    self._files_failed += 1
            except Exception as e:
                logger.warning(
                    "[%s] 处理文件失败 %s: %s",
                    self.name, file_info["name"], str(e)
                )
                self._files_failed += 1

        logger.info(
            "[%s] 转换完成: %d 成功, %d 失败 → %d 个文档",
            self.name, self._files_processed,
            self._files_failed, len(documents)
        )
        return documents

    async def _process_file(self, file_info: Dict[str, Any]) -> Optional[UnifiedDocument]:
        """处理单个文件"""
        file_path = Path(file_info["path"])
        ext = file_info["ext"]

        if ext in self._TEXT_EXTENSIONS:
            content = await self._read_text_file(file_path)
        elif ext == ".json":
            content = await self._read_json_file(file_path)
        elif ext == ".csv":
            content = await self._read_csv_file(file_path)
        elif ext in (".yaml", ".yml"):
            content = await self._read_text_file(file_path)
        elif ext == ".pdf":
            content = await self._read_pdf_file(file_path)
        elif ext == ".docx":
            content = await self._read_docx_file(file_path)
        elif ext == ".xlsx":
            content = await self._read_xlsx_file(file_path)
        else:
            # 未知格式：尝试作为文本读取
            try:
                content = await self._read_text_file(file_path)
            except UnicodeDecodeError:
                logger.warning(
                    "[%s] 无法读取文件（非文本或编码不支持）: %s",
                    self.name, file_path.name
                )
                return None

        if not content or len(content.strip()) < 5:
            logger.debug("[%s] 文件内容为空: %s", self.name, file_path.name)
            return None

        return UnifiedDocument(
            title=file_path.stem,
            content=content,
            source_type=SourceType.FILE,
            source_url=str(file_path.absolute()),
            metadata={
                "connector": self.name,
                "file_name": file_info["name"],
                "file_path": file_info["path"],
                "file_size": file_info["size"],
                "file_ext": ext,
                "mtime": file_info.get("mtime", ""),
            },
            language="zh",
        )

    async def _read_text_file(self, path: Path) -> str:
        """读取文本文件，自动检测编码"""
        # 先尝试 UTF-8
        try:
            return path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            pass

        # 尝试常见编码
        for encoding in ("utf-8-sig", "gbk", "gb2312", "latin-1"):
            try:
                return path.read_text(encoding=encoding)
            except (UnicodeDecodeError, LookupError):
                continue

        # 最后的回退
        return path.read_text(encoding="utf-8", errors="replace")

    async def _read_json_file(self, path: Path) -> str:
        """读取 JSON 文件，格式化为可读文本"""
        try:
            raw = path.read_text(encoding="utf-8")
            data = json.loads(raw)

            if isinstance(data, list):
                # JSON 数组：每条记录一行
                lines = []
                for i, item in enumerate(data):
                    if isinstance(item, dict):
                        lines.append(json.dumps(item, ensure_ascii=False))
                    else:
                        lines.append(str(item))
                return "\n".join(lines)
            else:
                return json.dumps(data, ensure_ascii=False, indent=2)

        except (json.JSONDecodeError, UnicodeDecodeError):
            return path.read_text(encoding="utf-8", errors="replace")

    async def _read_csv_file(self, path: Path) -> str:
        """读取 CSV 文件"""
        try:
            lines = []
            for encoding in ("utf-8", "utf-8-sig", "gbk"):
                try:
                    with open(path, "r", encoding=encoding, newline="") as f:
                        reader = csv.DictReader(f)
                        for row in reader:
                            lines.append(
                                " | ".join(
                                    f"{k}: {v}"
                                    for k, v in row.items()
                                )
                            )
                    break
                except (UnicodeDecodeError, csv.Error):
                    continue
            return "\n".join(lines[:10000])  # 限制最大行数
        except Exception as e:
            logger.warning("[%s] CSV 读取失败: %s", self.name, str(e))
            return path.read_text(encoding="utf-8", errors="replace")

    async def _read_pdf_file(self, path: Path) -> str:
        """读取 PDF 文件文本"""
        try:
            import PyPDF2
            text_parts = []
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            return "\n\n".join(text_parts)
        except ImportError:
            logger.warning(
                "[%s] 缺少 PyPDF2 依赖，跳过 PDF: %s",
                self.name, path.name
            )
            return ""
        except Exception as e:
            logger.warning(
                "[%s] PDF 读取失败 %s: %s",
                self.name, path.name, str(e)
            )
            return ""

    async def _read_docx_file(self, path: Path) -> str:
        """读取 DOCX 文件文本"""
        try:
            import docx
            doc = docx.Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n".join(paragraphs)
        except ImportError:
            logger.warning(
                "[%s] 缺少 python-docx 依赖，跳过 DOCX: %s",
                self.name, path.name
            )
            return ""
        except Exception as e:
            logger.warning(
                "[%s] DOCX 读取失败 %s: %s",
                self.name, path.name, str(e)
            )
            return ""

    async def _read_xlsx_file(self, path: Path) -> str:
        """读取 XLSX 文件文本"""
        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
            text_parts = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                text_parts.append(f"=== Sheet: {sheet_name} ===")
                for row in ws.iter_rows(values_only=True):
                    row_text = "\t".join(
                        str(cell) if cell is not None else ""
                        for cell in row
                    )
                    if row_text.strip():
                        text_parts.append(row_text)
            wb.close()
            return "\n".join(text_parts[:5000])
        except ImportError:
            logger.warning(
                "[%s] 缺少 openpyxl 依赖，跳过 XLSX: %s",
                self.name, path.name
            )
            return ""
        except Exception as e:
            logger.warning(
                "[%s] XLSX 读取失败 %s: %s",
                self.name, path.name, str(e)
            )
            return ""

    async def disconnect(self) -> None:
        """文件连接器断开（无持久连接需要关闭）"""
        self._set_status(ConnectorStatus.DISCONNECTED)
        logger.info("[%s] 已断开", self.name)

    @property
    def stats(self) -> Dict[str, Any]:
        base_stats = super().stats
        base_stats.update({
            "files_found": self._files_found,
            "files_processed": self._files_processed,
            "files_failed": self._files_failed,
        })
        return base_stats

    async def __aenter__(self):
        await self.connect()
        return self

    async def __aexit__(self, exc_type, exc_val, exc_tb):
        await self.disconnect()
