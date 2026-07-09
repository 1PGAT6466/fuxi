"""
pipeline.py — 少阳·消化 统一处理管线
合并胃(解析)+脾(存储)+肺(呼吸)+小肠(分类)的能力
"""
import asyncio
import logging
import struct
import time
from pathlib import Path
from typing import List, Dict
from dataclasses import dataclass, field

from src.models.chunk import Chunk
from src.models.event import Event
from src.models.entity import Entity
from src.infra.symbol_base import SymbolBase

logger = logging.getLogger("shaoyang.pipeline")


@dataclass
class PipelineResult:
    """管线处理结果"""
    source: str = ""
    file_path: str = ""
    raw_text: str = ""
    cleaned_text: str = ""
    tables: List[Dict] = field(default_factory=list)
    chunks: List[Chunk] = field(default_factory=list)
    events: List[Event] = field(default_factory=list)
    entities: List[Entity] = field(default_factory=list)
    metadata: Dict = field(default_factory=dict)
    skipped: bool = False
    duration_ms: float = 0
    errors: List[str] = field(default_factory=list)


class ShaoyangPipeline(SymbolBase):
    """少阳·消化 — 知识消化中枢"""

    def __init__(self, meridian):
        super().__init__(
            meridian=meridian,
            symbol_id="shaoyang",
            name="少阳·消化",
            emoji="🌱",
            description="知识消化中枢：文档进来 → 碎片 + 事件 + 实体出去"
        )
        self._processing = set()
        self._lock = asyncio.Lock()

    async def digest(self, file_path: str, source: str = "upload") -> PipelineResult:
        """统一处理入口"""
        start_time = time.time()
        self._set_status("processing")

        try:
            result = PipelineResult(source=source, file_path=file_path)

            # Step 1: 解析
            parsed = self._parse(file_path)
            result.raw_text = parsed.get("text", "")
            result.tables = parsed.get("tables", [])

            # Step 2: 清洗
            result.cleaned_text = self._clean(result.raw_text)

            # Step 3: 分块
            result.chunks = self._chunk(result.cleaned_text, result.tables)

            # Step 4: 分类
            for chunk in result.chunks:
                chunk.category = self._classify(chunk)

            # Step 5: 设置来源信息
            for chunk in result.chunks:
                chunk.file_hash = self._compute_hash(file_path)
                chunk.file_name = Path(file_path).name
                chunk.file_type = Path(file_path).suffix.lower()
                chunk.source_pipeline = source

            # Step 6: 向量化（如果 embedder 可用）
            await self._vectorize(result.chunks)

            # Step 7: 存储
            await self._save(result.chunks)

            # Step 8: Phase A — SAG 事件/实体提取
            await self._extract_events_entities(result)

            result.duration_ms = (time.time() - start_time) * 1000
            logger.info(f"[少阳] 完成: {file_path} → {len(result.chunks)} chunks, {len(result.events)} events, {len(result.entities)} entities, {result.duration_ms:.0f}ms")

            return result

        except Exception as e:  # TODO: Narrow exception type
            logger.error(f"[少阳] 处理失败: {file_path} → {e}")
            raise
        finally:
            self._set_status("idle")

    def _parse(self, file_path: str) -> Dict:
        """解析文件"""
        path = Path(file_path)
        ext = path.suffix.lower()

        try:
            if ext == ".pdf":
                return self._parse_pdf(file_path)
            elif ext in (".docx", ".doc"):
                return self._parse_docx(file_path)
            elif ext in (".xlsx", ".xls"):
                return self._parse_excel(file_path)
            elif ext == ".txt":
                return self._parse_text(file_path)
            elif ext == ".md":
                return self._parse_text(file_path)
            else:
                return self._parse_text(file_path)
        except Exception as e:  # TODO: Narrow exception type
            raise Exception(f"解析失败 ({ext}): {e}")

    def _parse_pdf(self, file_path: str) -> Dict:
        """PDF 解析"""
        try:
            import fitz
            doc = fitz.open(file_path)
            pages_text = []
            for page in doc:
                pages_text.append(page.get_text())
            doc.close()
            return {"text": "\n".join(pages_text), "tables": [], "metadata": {"parser": "fitz"}}
        except Exception:  # TODO: Narrow exception type
            return {"text": "", "tables": [], "metadata": {"parser": "none"}}

    def _parse_docx(self, file_path: str) -> Dict:
        """DOCX 解析"""
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return {"text": "\n".join(paragraphs), "tables": [], "metadata": {"parser": "python-docx"}}
        except Exception:  # TODO: Narrow exception type
            return {"text": "", "tables": [], "metadata": {"parser": "none"}}

    def _parse_excel(self, file_path: str) -> Dict:
        """Excel 解析"""
        try:
            import pandas as pd
            df = pd.read_excel(file_path)
            return {"text": df.to_string(index=False), "tables": [], "metadata": {"parser": "pandas"}}
        except Exception:  # TODO: Narrow exception type
            return {"text": "", "tables": [], "metadata": {"parser": "none"}}

    def _parse_text(self, file_path: str) -> Dict:
        """纯文本解析"""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return {"text": f.read(), "tables": [], "metadata": {"parser": "text"}}
        except Exception:  # TODO: Narrow exception type
            return {"text": "", "tables": [], "metadata": {"parser": "none"}}

    def _clean(self, text: str) -> str:
        """清洗文本"""
        import re
        text = re.sub(r'<[^>]+>', '', text)
        text = re.sub(r'https?://\S+', '', text)
        text = re.sub(r'\s+', ' ', text)
        return text.strip()

    def _chunk(self, text: str, tables: List[Dict]) -> List[Chunk]:
        """分块"""
        if not text or len(text) < 50:
            return [Chunk(text=text, chunk_index=0)] if text.strip() else []

        chunks = []
        chunk_size = 1000
        overlap = 100
        start = 0
        text_len = len(text)

        while start < text_len:
            end = min(start + chunk_size, text_len)
            if end < text_len:
                for sep in ['\n\n', '\n', '。', '；', '.', ';']:
                    last_sep = text.rfind(sep, start + chunk_size // 2, end)
                    if last_sep > start:
                        end = last_sep + len(sep)
                        break

            chunk_text = text[start:end].strip()
            if chunk_text and len(chunk_text) > 20:
                chunks.append(Chunk(text=chunk_text, chunk_index=len(chunks)))

            start = end - overlap if end < text_len else text_len

        for chunk in chunks:
            chunk.total_chunks = len(chunks)

        return chunks

    def _classify(self, chunk: Chunk) -> str:
        """分类"""
        try:
            from src.category_registry import match_category
            return match_category(chunk.text, file_ext=chunk.file_type, file_name=chunk.file_name) or "通用办公"
        except Exception:  # TODO: Narrow exception type
            return "通用办公"

    def _compute_hash(self, file_path: str) -> str:
        """计算文件哈希"""
        import hashlib
        with open(file_path, "rb") as f:
            return hashlib.sha256(f.read()).hexdigest()

    async def _vectorize(self, chunks: List[Chunk]):
        """向量化"""
        try:
            from src.db.vector_store import embed_texts
            embeddings = await embed_texts([c.text for c in chunks])
            if embeddings:
                for chunk, emb in zip(chunks, embeddings):
                    chunk.embedding = emb
        except Exception as e:  # TODO: Narrow exception type
            logger.warning(f"[少阳] 向量化失败: {e}")

    # FAKE-ASYNC: 本函数标记 async 仅为接口统一，内部同步执行
    async def _save(self, chunks: List[Chunk]):
        """存储：内存索引 + 向量库双写"""
        # 路1: BM25 全文索引 (memory_store)
        try:
            from src.db.memory_store import get_store
            store = get_store()
            chunk_dicts = [c.to_dict() for c in chunks]
            store.insert_many(chunk_dicts)
            logger.info(f"[少阳] BM25 存储 {len(chunk_dicts)} 个 chunk")
        except Exception as e:  # TODO: Narrow exception type
            logger.error(f"[少阳] BM25 存储失败: {e}")
            raise

        # 路2: ChromaDB 向量库 (vector_store) — v1.50 P0 修复
        try:
            from src.db.vector_store import get_vector_store
            vs = get_vector_store()
            if vs is None:
                logger.warning("[少阳] VectorStore 不可用，跳过向量写入")
                return
            # 收集已向量化的 chunk
            vectorized = [c for c in chunks if c.embedding is not None]
            if not vectorized:
                logger.info("[少阳] 无向量化数据，跳过向量写入")
                return
            ids = [f"{c.file_hash}_{c.chunk_index}" for c in vectorized]
            embeddings = [c.embedding for c in vectorized]
            metadatas = [
                {
                    k: str(v)[:512]
                    for k, v in c.to_dict().items()
                    if k not in ("text", "embedding")
                }
                for c in vectorized
            ]
            documents = [c.text for c in vectorized]
            success = vs.add(ids=ids, embeddings=embeddings, metadata=metadatas, documents=documents)
            if success:
                logger.info(f"[少阳] 向量存储 {len(vectorized)} 个 chunk → ChromaDB")
            else:
                logger.warning(f"[少阳] 向量写入失败 ({len(vectorized)} chunks)")
        except Exception as e:  # TODO: Narrow exception type
            logger.warning(f"[少阳] 向量写入异常（非致命）: {e}")

    # Phase A: SAG 事件/实体提取 + 写入 DB + 向量化
    async def _extract_events_entities(self, result: PipelineResult):
        """对每个 chunk 调用 SAGExtractor，将结果写入 events/entities 表

        LLM 不可用时的 fallback：跳过提取，标记事件为 'pending'。
        对提取成功的事件自动生成 embedding 并存入 events 表。
        """
        try:
            from src.shaoyang.extractor import SAGExtractor
            from src.db.memory_store import get_store
            from src.db.vector_store import embed_texts

            store = get_store()
            extractor = SAGExtractor()
            total_events = 0
            total_entities = 0

            # M-6: 收集实际提取的 Event/Entity 对象，不再丢失原始数据
            extracted_events: List[Dict] = []
            extracted_entities: List[Dict] = []

            for chunk in result.chunks:
                chunk_text = chunk.text
                if not chunk_text or len(chunk_text) < 50:
                    continue

                chunk_id = f"{chunk.file_hash}_{chunk.chunk_index}"

                try:
                    extraction = await extractor.extract(
                        chunk_text,
                        chunk_meta={
                            "chunk_id": chunk_id,
                            "file_hash": chunk.file_hash,
                            "file_name": chunk.file_name,
                            "category": chunk.category,
                        }
                    )

                    # Write events
                    for event_data in extraction.events:
                        event_data["chunk_id"] = chunk_id
                        event_data["file_hash"] = chunk.file_hash
                        event_data["file_name"] = chunk.file_name
                        event_data["event_type"] = event_data.get("action", "")
                        event_data.setdefault("entities", event_data.get("participants", []))
                        store.add_event(event_data)
                        total_events += 1
                        extracted_events.append(event_data)

                    # Write entities
                    for entity_data in extraction.entities:
                        entity_data.setdefault("chunk_ids", [chunk_id])
                        entity_data["file_hash"] = chunk.file_hash
                        entity_data["file_name"] = chunk.file_name
                        entity_data["source"] = "sag_extractor"
                        store.add_entity(entity_data)
                        total_entities += 1
                        extracted_entities.append(entity_data)

                except Exception as extract_err:  # TODO: Narrow exception type
                    error_msg = str(extract_err)
                    logger.warning(f"[少阳] SAG 提取失败 (chunk={chunk_id[:30]}...): {error_msg}")
                    store.add_event({
                        "event_id": f"evt_pending_{chunk_id[:60]}",
                        "chunk_id": chunk_id,
                        "title": "[PENDING] Extraction failed",
                        "content": f"SAG extraction pending. Error: {error_msg[:200]}",
                        "entities": [],
                        "event_type": "pending",
                        "file_hash": chunk.file_hash,
                        "file_name": chunk.file_name,
                        "status": "pending",
                    })
                    continue

            # Phase A Task 3: Event 向量化 — 对已入库事件生成 embedding
            if total_events > 0:
                await self._vectorize_events(total_events, store)

            # M-6: 保留实际提取的 events/entities 列表，不再用摘要对象覆盖
            result.events = [Event(
                event_id=ev.get("event_id", f"evt_{i}"),
                title=ev.get("title", ""),
            ) for i, ev in enumerate(extracted_events)] if extracted_events else []
            result.entities = [Entity(
                entity_id=ent.get("entity_id", f"ent_{i}"),
                name=ent.get("name", ""),
            ) for i, ent in enumerate(extracted_entities)] if extracted_entities else []

            if total_events > 0 or total_entities > 0:
                logger.info(f"[少阳] SAG 提取完成: {total_events} events, {total_entities} entities")
            else:
                logger.info("[少阳] SAG 提取: 无事件/实体产出（LLM 可能不可用）")

        except Exception as e:  # TODO: Narrow exception type
            logger.error(f"[少阳] SAG 提取管线异常: {e}")

    async def _vectorize_events(self, count: int, store):
        """Phase A Task 3: 为最近入库的活跃事件生成 embedding"""
        try:
            from src.db.vector_store import embed_texts

            # 获取最近入库且没有 embedding 的事件
            recent_events = store._db_conn.execute(
                "SELECT id, event_id, content FROM events WHERE status='active' AND embedding IS NULL ORDER BY id DESC LIMIT ?",
                (count,)
            ).fetchall()

            if not recent_events:
                return

            contents = [row[2] or row[1] for row in recent_events]  # content or fallback to event_id
            embeddings = await embed_texts(contents)

            if embeddings:
                for (row_id, _, _), emb in zip(recent_events, embeddings):
                    if emb:
                        blob = struct.pack(f'{len(emb)}f', *emb)
                        store._db_conn.execute(
                            "UPDATE events SET embedding=? WHERE id=?",
                            (blob, row_id)
                        )
                store._db_conn.commit()
                logger.info(f"[少阳] Event 向量化完成: {len(recent_events)} 条")
            else:
                logger.info("[少阳] Event 向量化跳过 (embedder 不可用)")

        except Exception as e:  # TODO: Narrow exception type
            logger.warning(f"[少阳] Event 向量化失败: {e}")


# 全局实例
